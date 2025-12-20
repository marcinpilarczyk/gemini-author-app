import streamlit as st
import google.generativeai as genai
from google.generativeai import caching
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import datetime
import re
import sqlite3
import json
import os
from docx import Document
from io import BytesIO
import time

# --- PAGE CONFIGURATION ---
st.set_page_config(page_title="Gemini 3 Author Studio", layout="wide")
st.title("Drafting with Gemini 3 Pro (Full Studio Edition)")

# --- DATABASE SETUP ---
DB_NAME = "my_novel.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS books (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT DEFAULT 'Untitled Book',
                    concept TEXT,
                    outline TEXT
                )''')
    c.execute('''CREATE TABLE IF NOT EXISTS chapters (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    book_id INTEGER,
                    chapter_num INTEGER,
                    content TEXT,
                    summary TEXT,
                    FOREIGN KEY(book_id) REFERENCES books(id)
                )''')
    conn.commit()
    conn.close()

def get_all_books():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT id, title FROM books ORDER BY id")
    books = c.fetchall()
    conn.close()
    return books

def create_new_book(title):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT INTO books (title, concept, outline) VALUES (?, '', '')", (title,))
    new_id = c.lastrowid
    conn.commit()
    conn.close()
    return new_id

def load_active_book(book_id):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM books WHERE id=?", (book_id,))
    book = c.fetchone()
    c.execute("SELECT * FROM chapters WHERE book_id=? ORDER BY chapter_num", (book_id,))
    chapters = c.fetchall()
    conn.close()
    return book, chapters

def get_chapters(book_id):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM chapters WHERE book_id=? ORDER BY chapter_num ASC", (book_id,))
    chapters = c.fetchall()
    conn.close()
    return chapters

def update_book_meta(book_id, title, concept, outline):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("UPDATE books SET title=?, concept=?, outline=? WHERE id=?", (title, concept, outline, book_id))
    conn.commit()
    conn.close()

def save_chapter(book_id, num, content, summary=""):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT id, summary FROM chapters WHERE book_id=? AND chapter_num=?", (book_id, num))
    existing = c.fetchone()
    if existing:
        current_sum = summary if summary else (existing[1] if existing[1] else "")
        c.execute("UPDATE chapters SET content=?, summary=? WHERE id=?", (content, current_sum, existing[0]))
    else:
        c.execute("INSERT INTO chapters (book_id, chapter_num, content, summary) VALUES (?, ?, ?, ?)", 
                  (book_id, num, content, summary))
    conn.commit()
    conn.close()

def delete_last_chapter(book_id, num):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("DELETE FROM chapters WHERE book_id=? AND chapter_num=?", (book_id, num))
    conn.commit()
    conn.close()

def reset_db():
    if os.path.exists(DB_NAME):
        os.remove(DB_NAME)
    init_db()

init_db()

# --- MODEL CONFIG ---
MODEL_NAME = "gemini-3-pro-preview" 

safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# --- HELPERS ---
def generate_summary(chapter_text):
    if not chapter_text or len(chapter_text.strip()) < 50: return ""
    prompt = f"""Analyze the following chapter and provide a technical summary for an author's continuity ledger.
    
    Output Format:
    1. Narrative Summary: A concise paragraph of what actually happened (the events and plot movements).
    2. Facts/Items/Injuries: Key details (character descriptions, specific items found/used, new wounds, locations).
    3. Pacing: Analysis of the scene's intensity shifts (Start, Middle, End).
    
    Chapter Text:
    {chapter_text[:12000]}"""
    
    try:
        model = genai.GenerativeModel(MODEL_NAME, safety_settings=safety_settings)
        return model.generate_content(prompt).text
    except Exception as e: return f"Error: {e}"

def normalize_text(text, mode="standard"):
    if not text: return ""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    paragraphs = re.split(r'\n\s*\n', text)
    clean_paragraphs = [p.strip() for p in paragraphs if p.strip()]
    if mode == "tight": return '\n'.join(clean_paragraphs)
    else: return '\n\n'.join(clean_paragraphs)

def create_docx(full_text, title):
    doc = Document()
    doc.add_heading(title, 0)
    normalized = normalize_text(full_text, mode="standard")
    paragraphs = normalized.split('\n\n')
    for p_text in paragraphs:
        if not p_text.strip(): continue
        if p_text.startswith("## Chapter"):
            doc.add_heading(p_text.replace("## ", "").strip(), level=1)
        elif p_text.startswith("## "):
            doc.add_heading(p_text.replace("## ", "").strip(), level=2)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', p_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    run = p.add_run(part[2:-2]); run.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    run = p.add_run(part[1:-1]); run.italic = True
                else: p.add_run(part)
    return doc

def get_or_create_cache(bible_text, outline_text):
    static_content = f"### BIBLE\n{bible_text}\n\n### OUTLINE\n{outline_text}"
    if 'cache_name' in st.session_state:
        try:
            cache = genai.caching.CachedContent.get(name=st.session_state.cache_name)
            cache.update(ttl=datetime.timedelta(hours=2))
            return cache.name
        except: del st.session_state.cache_name
    try:
        cache = genai.caching.CachedContent.create(
            model=MODEL_NAME, display_name="book_bible_v1", contents=[static_content], ttl=datetime.timedelta(hours=2)
        )
        st.session_state.cache_name = cache.name
        return cache.name
    except: return None

# --- SIDEBAR ---
with st.sidebar:
    st.header("🔑 Settings")
    if "GOOGLE_API_KEY" in st.secrets: api_key = st.secrets["GOOGLE_API_KEY"]
    else: api_key = st.text_input("Enter Google API Key", type="password")
    
    available_models = ["gemini-3-pro-preview", "gemini-3-flash-preview", "gemini-2.0-flash-exp", "gemini-1.5-pro-latest"]
    if "model_name" not in st.session_state: st.session_state.model_name = available_models[0]
    selected_model = st.selectbox("🤖 Engine", available_models, index=available_models.index(st.session_state.model_name))
    if selected_model != st.session_state.model_name:
        st.session_state.model_name = selected_model
        st.session_state.cache_name = None; st.rerun()
    MODEL_NAME = st.session_state.model_name
    
    st.divider()
    st.subheader("📚 Library")
    all_books = get_all_books()
    if not all_books:
        first_id = create_new_book("My First Book"); st.session_state.active_book_id = first_id; st.rerun()
    
    if "active_book_id" not in st.session_state:
        st.session_state.active_book_id = all_books[0]['id']
    
    book_opts = {b['id']: b['title'] for b in all_books}
    try:
        current_book_index = list(book_opts.keys()).index(st.session_state.active_book_id)
    except ValueError:
        current_book_index = 0
        
    sel_id = st.selectbox("Current Book", options=book_opts.keys(), format_func=lambda x: book_opts[x], index=current_book_index)
    if sel_id != st.session_state.active_book_id:
        st.session_state.active_book_id = sel_id; st.session_state.cache_name = None; st.rerun()

    with st.popover("➕ New Book"):
        nt = st.text_input("Title", "Untitled")
        if st.button("Create"):
            nid = create_new_book(nt)
            st.session_state.active_book_id = nid
            st.rerun()

    st.divider()
    
    with st.expander("💾 Backup & Restore"):
        st.caption("Since the server is temporary, download your database to save your work permanently.")
        if os.path.exists(DB_NAME):
            with open(DB_NAME, "rb") as f:
                st.download_button("📥 Download Database (.db)", f, file_name=f"author_studio_backup_{datetime.date.today()}.db")
        
        st.divider()
        uploaded_db = st.file_uploader("📤 Restore from Backup", type="db")
        if uploaded_db:
            if st.button("Overwrite Current with Backup"):
                with open(DB_NAME, "wb") as f:
                    f.write(uploaded_db.getbuffer())
                st.success("Project Restored! Reloading...")
                time.sleep(1)
                st.rerun()

    with st.expander("⚠️ Import Manuscript"):
        imp_txt = st.text_area("Paste Full Text (Will split by 'Chapter X')", height=200)
        if st.button("Import"):
            if imp_txt:
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("DELETE FROM chapters WHERE book_id=?", (st.session_state.active_book_id,))
                chunks = re.split(r'(?i)(chapter\s+\d+)', imp_txt)
                cn, cc = 0, ""
                for ch in chunks:
                    if re.match(r'(?i)chapter\s+\d+', ch.strip()):
                        if cn > 0:
                            cl = normalize_text(cc)
                            if cl: c.execute("INSERT INTO chapters (book_id, chapter_num, content, summary) VALUES (?, ?, ?, ?)", (st.session_state.active_book_id, cn, cl, ""))
                        cn += 1
                        cc = ""
                    else: cc += ch
                if cn > 0:
                    cl = normalize_text(cc)
                    if cl: c.execute("INSERT INTO chapters (book_id, chapter_num, content, summary) VALUES (?, ?, ?, ?)", (st.session_state.active_book_id, cn, cl, ""))
                conn.commit()
                conn.close()
                st.success("Imported!")
                st.rerun()

    with st.expander("⚡ Memory Management"):
        overwrite_summaries = st.checkbox("Overwrite existing summaries", value=False)
        if st.button("Process Summaries"):
            if not api_key: st.error("Need Key")
            else:
                genai.configure(api_key=api_key)
                conn = sqlite3.connect(DB_NAME); conn.row_factory = sqlite3.Row
                c = conn.cursor()
                c.execute("SELECT * FROM chapters WHERE book_id=? AND content IS NOT NULL", (st.session_state.active_book_id,))
                rows = c.fetchall()
                if not rows: st.warning("No chapters found.")
                else:
                    bar = st.progress(0); status = st.empty()
                    for i, r in enumerate(rows):
                        if not r['summary'] or len(r['summary']) < 10 or overwrite_summaries:
                            status.text(f"Summarizing Ch {r['chapter_num']}...")
                            s = generate_summary(r['content'])
                            if s and not s.startswith("Error"):
                                c2 = conn.cursor()
                                c2.execute("UPDATE chapters SET summary=? WHERE id=?", (s, r['id']))
                                conn.commit()
                        bar.progress((i+1)/len(rows))
                    status.text("Done."); st.success("Backfill Complete!"); st.rerun()

    if st.button("🔴 Reset Database"):
        reset_db(); st.session_state.clear(); st.rerun()

# --- MAIN LOGIC ---
if not api_key: st.warning("👈 Enter API Key"); st.stop()
genai.configure(api_key=api_key)
model = genai.GenerativeModel(MODEL_NAME, safety_settings=safety_settings)

active_book, chapter_data = load_active_book(st.session_state.active_book_id)
current_title = active_book['title']
current_concept = active_book['concept']
current_outline = active_book['outline']

full_text = ""
rolling_sum = ""
existing_chapters = {}
history_list = []

for r in chapter_data:
    history_list.append(r)
    existing_chapters[r['chapter_num']] = r['content']
    full_text += f"\n\n## Chapter {r['chapter_num']}\n\n{r['content']}"
    if r['summary']: rolling_sum += f"\n\n**Ch {r['chapter_num']}:**\n{r['summary']}"

st.subheader(f"📖 {current_title}")
t1, t2, t3, t4, t5 = st.tabs(["1. Bible", "2. Writer", "3. Manuscript", "4. Publisher", "5. Editor"])

# TAB 1: BIBLE
with t1:
    c1, c2 = st.columns(2)
    with c1: nti = st.text_input("Title", value=current_title); nc = st.text_area("Concept", value=current_concept, height=500)
    with c2: st.write(""); st.write(""); no = st.text_area("Outline", value=current_outline, height=500)
    if nc!=current_concept or no!=current_outline or nti!=current_title:
        if st.button("💾 Save Bible"): update_book_meta(st.session_state.active_book_id, nti, nc, no); st.rerun()

# TAB 2: WRITER
with t2:
    if "selected_chap" not in st.session_state: st.session_state.selected_chap = len(history_list) + 1
    if "editor_mode" not in st.session_state: st.session_state.editor_mode = False
    
    c_sel1, c_sel2 = st.columns([1, 4])
    with c_sel1:
        chap_num = st.number_input("Chapter #", min_value=1, value=st.session_state.selected_chap, step=1)
        st.session_state.selected_chap = chap_num
    with c_sel2:
        st.write(""); st.write("")
        if chap_num in existing_chapters and not st.session_state.editor_mode:
            if st.button(f"✏️ Load Chapter {chap_num} for Editing"):
                st.session_state.ed_con = existing_chapters[chap_num]; st.session_state.editor_mode = True; st.rerun()
    
    st.divider()
    if st.button(f"🔮 Auto-Fetch Plan for Ch {chap_num}"):
        with st.spinner("Fetching..."):
            p = f"Access Outline. Copy section for **Chapter {chap_num}** VERBATIM."
            try:
                cn = get_or_create_cache(nc, no)
                res = genai.GenerativeModel.from_cached_content(cached_content=genai.caching.CachedContent.get(name=cn)).generate_content(p) if cn else model.generate_content(f"{no}\n\n{p}")
                st.session_state[f"pl_{chap_num}"] = res.text; st.rerun()
            except Exception as e: st.error(f"Error: {e}")
    
    cp = st.session_state.get(f"pl_{chap_num}", "")
    ci = st.text_area("Chapter Plan / Instructions", value=cp, height=150)

    if not st.session_state.editor_mode:
        btn_label = f"🚀 Write Chapter {chap_num}" if chap_num not in existing_chapters else f"🔄 Re-Write Chapter {chap_num}"
        if st.button(btn_label, type="primary"):
            with st.spinner("Writing..."):
                cn = get_or_create_cache(nc, no)
                prev_text = existing_chapters.get(chap_num - 1, "")[-3000:] if chap_num > 1 else ""
                dp = f"### CONTEXT\n{rolling_sum}\n### PREV TEXT\n...{prev_text}\n### PLAN\n{ci}\n### TASK\nWrite Ch {chap_num}. Use Markdown headers."
                try:
                    res = genai.GenerativeModel.from_cached_content(cached_content=genai.caching.CachedContent.get(name=cn), safety_settings=safety_settings).generate_content(dp) if cn else model.generate_content(f"{nc}\n{no}\n{dp}")
                    st.session_state.ed_con = normalize_text(res.text); st.session_state.editor_mode = True; st.rerun()
                except Exception as e: st.error(f"Error: {e}")
    else:
        st.info(f"📝 Editing Chapter {chap_num}")
        tab_edit, tab_prev = st.tabs(["✍️ Edit", "👁️ Preview"])
        with tab_edit: et = st.text_area("Content", value=st.session_state.ed_con, height=600, key="ed_con_ta")
        with tab_prev: st.markdown(et)
        c1, c2 = st.columns([1,4])
        with c1:
            if st.button("💾 Save"):
                with st.spinner("Saving..."):
                    sm = generate_summary(et); save_chapter(st.session_state.active_book_id, chap_num, et, sm)
                    st.session_state.editor_mode = False; del st.session_state.ed_con; st.rerun()
        with c2:
            if st.button("❌ Discard"):
                st.session_state.editor_mode = False; del st.session_state.ed_con; st.rerun()

    if not st.session_state.editor_mode:
        st.divider()
        prev_chap_idx = chap_num - 1
        if prev_chap_idx in existing_chapters:
            prev_summary = next((r['summary'] for r in history_list if r['chapter_num'] == prev_chap_idx), "No summary.")
            with st.expander(f"⬅️ Reference: Chapter {prev_chap_idx} (Previous)"):
                st.info(prev_summary); st.markdown(existing_chapters[prev_chap_idx])
        
        if history_list:
            with st.expander("📚 View All Saved Chapters"):
                if st.button("Undo Last Chapter Addition"):
                    delete_last_chapter(st.session_state.active_book_id, history_list[-1]['chapter_num']); st.rerun()
                for h in reversed(history_list):
                    with st.expander(f"Ch {h['chapter_num']} View"):
                        st.info(h['summary']); st.markdown(h['content'])

# TAB 3: MANUSCRIPT
with t3:
    if st.button("📄 Export Word"):
        d = create_docx(full_text, current_title); b = BytesIO(); d.save(b); b.seek(0)
        st.download_button("Download", b, f"{current_title}.docx")
    mt1, mt2 = st.tabs(["📖 Reading View", "📝 Raw Text"])
    with mt1: st.markdown(full_text)
    with mt2: st.text_area("Manuscript", value=full_text, height=600)

# TAB 4: PUBLISHER
with t4:
    if st.button("🧬 Analyze DNA"):
        with st.spinner("Analyzing..."):
            try:
                res = model.generate_content(f"Analyze for KDP:\n{nc}\n{no}\n{rolling_sum}\nReturn: GENRE, TROPES, TONE").text
                st.session_state.dna_res = res; st.rerun()
            except Exception as e: st.error(f"Error: {e}")
    if "dna_res" in st.session_state: st.info(st.session_state.dna_res)

# TAB 5: EDITOR
with t5:
    st.header("🧐 Smart Consistency Editor")
    def apply_minimal_fix(chap_num, old_text, new_text):
        conn = sqlite3.connect(DB_NAME); c = conn.cursor()
        c.execute("SELECT content FROM chapters WHERE book_id=? AND chapter_num=?", (st.session_state.active_book_id, chap_num))
        row = c.fetchone()
        if row:
            updated = row[0].replace(old_text.strip(), new_text.strip())
            if updated != row[0]:
                ns = generate_summary(updated)
                c.execute("UPDATE chapters SET content=?, summary=? WHERE book_id=? AND chapter_num=?", (updated, ns, st.session_state.active_book_id, chap_num))
                conn.commit(); st.success(f"Fixed Ch {chap_num}!"); time.sleep(1)
            else: st.warning("Exact match not found.")
        conn.close()

    strict_config = genai.types.GenerationConfig(temperature=0.1, top_p=0.95, max_output_tokens=65000)
    if st.button("🔍 Run Full Logic Scan"):
        if len(full_text) < 500: st.error("Too short.")
        else:
            with st.spinner("Analyzing..."):
                prompt = f"""You are a Continuity Editor. Identify logic breaks and propose MINIMAL FIXES.
                ### THE MANUSCRIPT
                {full_text}
                
                OUTPUT FORMAT:
                [Narrative Report]
                ---FIX_BLOCK---
                [ {{"chapter": 1, "find": "old text", "replace": "new text"}} ]
                ---END_FIX_BLOCK---
                """
                try:
                    cn = get_or_create_cache(nc, no)
                    response = genai.GenerativeModel.from_cached_content(cached_content=genai.caching.CachedContent.get(name=cn)).generate_content(prompt, generation_config=strict_config) if cn else model.generate_content(prompt, generation_config=strict_config)
                    if hasattr(response, 'text') and response.text:
                        st.session_state.editor_report = response.text
                        try:
                            st.session_state.parsed_fixes = json.loads(response.text.split("---FIX_BLOCK---")[1].split("---END_FIX_BLOCK---")[0])
                        except: st.session_state.parsed_fixes = []
                        st.rerun()
                except Exception as e: st.error(f"Error: {e}")

    if "editor_report" in st.session_state:
        st.markdown(st.session_state.editor_report.split("---FIX_BLOCK---")[0])
        if st.session_state.get("parsed_fixes"):
            st.divider(); st.subheader("🛠️ Propose Fixes")
            for i, fix in enumerate(st.session_state.parsed_fixes):
                with st.expander(f"Ch {fix['chapter']} Suggestion"):
                    st.write(f"**Find:** {fix['find']}"); st.write(f"**Replace:** {fix['replace']}")
                    if st.button("Apply", key=f"app_{fix['chapter']}_{i}"):
                        apply_minimal_fix(fix['chapter'], fix['find'], fix['replace'])
                        st.session_state.parsed_fixes.pop(i); st.rerun()
